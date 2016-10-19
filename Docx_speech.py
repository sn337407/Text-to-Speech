import pyttsx
import docx
engine=pyttsx.init()
doc = docx.Document('demo.docx')
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
read=getText('demo.docx')
engine.say(read)
engine.runAndWait()
print('Success')
      
                       